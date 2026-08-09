# -*- coding: utf-8 -*-
r"""
通用 Markdown → Word(.docx) 转换脚本
=====================================
支持: #/##/### 标题、- 列表、| 表格、> 引用、``` 代码块、**加粗**、--- 分隔线。
用法: python make_docx.py 输入.md [输出.docx]
      (省略输出路径则生成同名 .docx)
"""
import os
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor


def add_rich(p, text):
    """把 **加粗** 片段写入段落 p"""
    parts = text.split("**")
    for pi, part in enumerate(parts):
        r = p.add_run(part)
        if pi % 2 == 1:
            r.bold = True


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else "platform_log.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(src)[0] + ".docx"

    doc = Document()
    in_code = False
    lines = open(src, encoding="utf-8").read().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # 代码块
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            i += 1
            continue
        # 标题
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 4)
            doc.add_heading(line.lstrip("# ").strip(), level)
        # 表格
        elif line.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            if len(rows) > 1 and all(re.fullmatch(r":?-+:?", c) for c in rows[1]):
                rows.pop(1)
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci in range(cols):
                    table.rows[ri].cells[ci].text = row[ci] if ci < len(row) else ""
            doc.add_paragraph()
            continue
        # 列表
        elif line.startswith("- "):
            add_rich(doc.add_paragraph(style="List Bullet"), line[2:])
        # 引用
        elif line.startswith("> "):
            p = doc.add_paragraph()
            add_rich(p, line[2:])
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        # 分隔线 / 空行
        elif line.strip() in ("", "---", "***"):
            pass
        else:
            add_rich(doc.add_paragraph(), line)
        i += 1

    doc.save(dst)
    print("已生成:", dst)


if __name__ == "__main__":
    main()
